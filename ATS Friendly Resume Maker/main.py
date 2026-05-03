"""
ATS Friendly - Resume Builder Backend — FastAPI + Gemini

Receives:
  - Past resume uploads (PDF / DOCX / TXT) — extracts plain text and
    auto-parses it into the CV JSON schema using the Gemini SDK
    (google-genai) with structured output.
  - Target job description text
  - Current CV form data (JSON matching cv_template.json)

Stores them in an in-memory session keyed by session_id.

Configure:
  Set GEMINI_API_KEY in environment (or in a .env file in this folder).
  Get a key at https://aistudio.google.com/apikey

Run:
  uvicorn main:app --reload --port 8000
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
import io
import os
import re
import json
import uuid
import copy

from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pathlib import Path

def _static_dir() -> Path:
    import sys
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS)        
    return Path(__file__).parent         



# ---- Optional: load .env if python-dotenv is installed ----------------
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ---- Optional file-parsing dependencies -------------------------------
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ---- Gemini SDK -------------------------------------------------------
try:
    from google import genai
    from google.genai import types as genai_types
    HAS_GENAI = True
except ImportError:
    HAS_GENAI = False



# =====================================================================
# App + CORS
# =====================================================================

app = FastAPI(
    title="Resume Builder API",
    description="Backend for the resume builder UI. Uses Gemini to parse "
                "uploaded resumes into the CV JSON schema.",
    version="0.2.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =====================================================================
# Pydantic models — mirror cv_template.json exactly
# These are used both for FastAPI request validation AND as the
# response_schema for Gemini structured output.
# =====================================================================

class Personal(BaseModel):
    name: str = ""
    title: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    website: str = ""
    websiteUrl: str = ""    
    linkedin: str = ""
    linkedinUrl: str = ""   
    github: str = ""
    githubUrl: str = ""     


class Experience(BaseModel):
    id: int = 0
    company: str = ""
    role: str = ""
    location: str = ""
    startDate: str = ""
    endDate: str = ""
    bullets: List[str] = Field(default_factory=list)


class Education(BaseModel):
    id: int = 0
    school: str = ""
    degree: str = ""
    location: str = ""
    startDate: str = ""
    endDate: str = ""
    coursework: str = ""   # e.g. "Cloud Computing, Software Engineering, Algorithms"
    gpa: str = ""          # e.g. "3.89/4.0"
    details: str = ""      # optional free-form: honors / awards / thesis


class Skill(BaseModel):
    id: int = 0
    category: str = ""
    items: str = ""


class Project(BaseModel):
    id: int = 0
    name: str = ""
    tech: str = ""
    link: str = ""
    bullets: List[str] = Field(default_factory=list)


class Certification(BaseModel):
    id: int = 0
    name: str = ""
    issuer: str = ""
    date: str = ""


class CVData(BaseModel):
    personal: Personal = Field(default_factory=Personal)
    summary: str = ""
    experience: List[Experience] = Field(default_factory=list)
    education: List[Education] = Field(default_factory=list)
    skills: List[Skill] = Field(default_factory=list)
    projects: List[Project] = Field(default_factory=list)
    certifications: List[Certification] = Field(default_factory=list)


class JDPayload(BaseModel):
    session_id: str = "default"
    text: str


class CVPayload(BaseModel):
    session_id: str = "default"
    data: CVData


# ----- New: JD keyword extraction & tailoring -------------------------

class JDKeyword(BaseModel):
    """A single keyword/phrase pulled from the job description."""
    term: str = ""
    # "skill" | "tool" | "concept" | "responsibility" | "qualification"
    category: str = "skill"
    # "must" | "nice"
    importance: str = "must"


class JDKeywordList(BaseModel):
    keywords: List[JDKeyword] = Field(default_factory=list)


class TailorPayload(BaseModel):
    session_id: str = "default"


class ScorePayload(BaseModel):
    session_id: str = "default"
    # Optional: score against a CV passed in the body (for live preview while
    # the user is typing) instead of the CV stored on the session.
    data: Optional[CVData] = None


class UpdateKeywordsPayload(BaseModel):
    """Replace the session's JD keyword list. Used when the user manually
    removes (or, eventually, edits) keywords from the ATS panel.
    """
    session_id: str = "default"
    keywords: List[JDKeyword] = Field(default_factory=list)
    # Optional: tailored CV held on the frontend. When present, we score it
    # alongside the original so the panel's versus row stays in sync.
    tailored_data: Optional[CVData] = None


class TailorEdit(BaseModel):
    """One change Gemini made — used to compute highlight spans server-side."""
    path: str = ""          # e.g. "experience.0.bullets.2", "skills.1.items"
    new_text: str = ""      # the replacement text
    inserted_terms: List[str] = Field(default_factory=list)  # JD terms woven in


class TailorResult(BaseModel):
    """The full structured response we ask Gemini to return."""
    edits: List[TailorEdit] = Field(default_factory=list)
    gaps: List[JDKeyword] = Field(default_factory=list)
    summary: str = ""


# =====================================================================
# Gemini client
# =====================================================================

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
print(f"[gemini] API key {'found' if GEMINI_API_KEY else 'not found'} in environment")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash").strip() or "gemini-2.5-flash"

gemini_client: Optional["genai.Client"] = None
if HAS_GENAI and GEMINI_API_KEY:
    try:
        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        print(f"[gemini] client initialized · model={GEMINI_MODEL}")
    except Exception as e:
        print(f"[gemini] init failed: {e}")
        gemini_client = None
elif HAS_GENAI:
    print("[gemini] SDK installed but GEMINI_API_KEY not set — auto-parse disabled")
else:
    print("[gemini] google-genai not installed — auto-parse disabled")


PARSE_PROMPT = """You are an expert resume parser. Extract structured information \
from the resume text below into the provided JSON schema.

CRITICAL RULES:
1. Use ONLY what's in the resume — never invent, infer, or hallucinate.
2. For missing fields, use empty strings ("") or empty arrays ([]).
3. Use sequential integer IDs starting from 1 within each array \
(experience, projects, education, skills, certifications).
4. OCR cleanup: fix obvious errors like "Al" → "AI", "lO" → "10", \
"Silerio" → "Silero", but never alter intended meaning.
5. URLs: strip "https://" / "http://" prefixes (e.g. "github.com/user", \
not "https://github.com/user"). Do NOT add www.
6. Dates: prefer "MMM YYYY" format (e.g. "Mar 2022", "Aug 2016"). \
Use "Present" for current roles. Year-only is acceptable if that's all the source gives. \
For education with "Expected YYYY", put that string in endDate.
7. Bullets (experience): each is a complete sentence ending with a period. Preserve original \
wording wherever possible — clean up only obvious typos/OCR.
8. Skills: organize into clear categories such as "Programming Languages", \
"Frameworks", "Cloud", "Databases", "DevOps & Tools", "AI/ML Tools". \
Items in each category as a comma-separated string.
9. Projects: this template renders ONE bullet per project (the project name in bold \
followed by an inline description). Therefore, put the full project description into \
`bullets` as a SINGLE element list — i.e. `bullets: ["<one paragraph description>"]`. \
Extract tech stack into the dedicated `tech` field (comma-separated) only if explicitly \
called out separately; otherwise leave `tech` empty and keep tech mentions inside the \
description bullet.
10. Personal `title`: most recent role, or the title shown at the top of the resume. \
If the resume header has only the name and contact line (no title underneath), leave `title` empty.
11. If a "Tech:" preamble appears in an experience entry, do NOT include it as a \
bullet — that information lives in `skills` and is implicit in achievements.
12. Certifications: if a single line lists multiple certs (e.g. \
"Microsoft Certified: Azure AI Engineer Associate, Developer Associate, AI Fundamentals"), \
split them into separate entries each with the appropriate issuer. \
Note: if certifications appear inside the SKILLS section as a bullet (e.g. \
"Certifications and MOOCs: AWS Certified Cloud Practitioner, ..."), keep them in `skills` \
under that category — do NOT also duplicate into the `certifications` array.
13. Education `coursework`: the comma-separated list of courses listed under \
"Relevant Coursework" or similar — strip the leading label and store the bare list.
14. Education `gpa`: the GPA value as it appears on the resume, including the scale \
(e.g. "3.89/4.0", "9.06/10"). Do NOT include the "GPA:" label.
15. Education `details`: use only for content that doesn't fit `coursework` or `gpa` \
(e.g. honors, thesis title, study-abroad note). Leave empty when the resume has \
just coursework + GPA.
16. Personal contact URLs: when the resume contains a clickable link or full URL for \
LinkedIn, GitHub, or a personal website/portfolio, put the FULL https://... URL into \
`linkedinUrl`, `githubUrl`, and `websiteUrl` respectively. Use friendly short labels \
("LinkedIn", "Github", "Portfolio Website") for the corresponding `linkedin`, `github`, \
and `website` fields. If only a URL is present and no label, mirror the URL into the \
label field too. If neither is present, leave both empty.

RESUME TEXT:
\"\"\"
{resume_text}
\"\"\"
"""


async def parse_resume_with_gemini(resume_text: str) -> Optional[CVData]:
    """
    Send resume text to Gemini with the CVData schema as response_schema.
    Returns a parsed CVData object, or None on failure.
    """
    if not gemini_client or not resume_text.strip():
        return None

    prompt = PARSE_PROMPT.format(resume_text=resume_text)

    try:
        response = await gemini_client.aio.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=CVData,
                temperature=0.1,  # low for deterministic extraction
            ),
        )
        # response.parsed is the typed CVData when response_schema is a Pydantic model
        parsed = response.parsed
        if isinstance(parsed, CVData):
            return _normalize_ids(parsed)
        # Fallback — parse the JSON text manually
        if response.text:
            return _normalize_ids(CVData.model_validate_json(response.text))
        return None
    except Exception as e:
        print(f"[gemini] parse failed: {e}")
        raise


def _normalize_ids(cv: CVData) -> CVData:
    """Make sure IDs are sequential integers starting at 1 within each list."""
    for items in (cv.experience, cv.education, cv.skills, cv.projects, cv.certifications):
        for i, item in enumerate(items, start=1):
            item.id = i
    return cv


# =====================================================================
# JD keyword extraction (Gemini, strict no-hallucination prompt)
# =====================================================================

JD_KEYWORD_PROMPT = """You extract ATS keywords from the job description \
below. Recruiters and ATS systems search for specific phrases — your output \
must reflect what they would search for, not a fragmented word list and \
not a bloated one either.

RULES IN PRIORITY ORDER:

1. EXTRACT ONLY. Every keyword must literally appear in the JD (verbatim or \
   trivial inflection like singular/plural). Never invent or infer. Verify \
   each keyword appears in the JD before returning.

2. CAPTURE THE WHOLE PHRASE, NOT THE PIECES. When the JD uses a meaningful \
   multi-word term, extract the term — not its parts. This is the most \
   common failure mode and the most important rule.
   - "build automation solutions"   → "build automation solutions"
                                       NOT "build" + "automation" + "solutions"
   - "LLM integration"              → "LLM integration"
   - "data pipeline development"    → "data pipeline development"
   - "REST API design"              → "REST API design"
   - "performance optimization"     → "performance optimization"
   Atomic words are valid when the JD uses them alone: "Python", \
   "Snowflake", "Kubernetes", "AWS", "Terraform" — keep as-is.

3. USE NATURAL INDUSTRY PHRASING. Keywords must read the way they appear in \
   real resumes and job descriptions. Do NOT invent unnatural noun-forms by \
   gerundizing every verb.
   - "identify automation opportunities" — keep as-is, do NOT convert to \
     "automation opportunity identification".
   - "prompt engineering" — keep as-is, do NOT inflate to \
     "prompt engineering strategy building".
   - "task management" — yes; "manage tasks effectively" — no.
   - "API development" — yes; "API development design implementation" — no.
   If a phrase sounds like something you'd never see on LinkedIn, it's wrong.

4. ONE FORM PER CONCEPT — NO PARALLEL VARIANTS. If you've already captured \
   a concept, do NOT also include design/build/maintain/develop variants of \
   the same thing.
   - Pick ONE: "automation solutions" OR "build automation solutions". Not both.
   - Pick ONE: "LLM-powered workflows" OR "LLM workflow design" OR "LLM \
     workflow implementation". Not three.
   - Pick ONE: "automation script development" OR "automation script \
     maintenance" OR "automation scripts". The unifying noun phrase wins.
   When in doubt, prefer the noun phrase form over the verb-led form, and \
   prefer the shortest version that retains meaning.

5. CAPTURE COMMON COMBINED FORMS WHEN THE JD SUPPORTS THEM. These specific \
   action+skill pairings are high-value ATS phrases. Include them when the \
   JD's content genuinely covers the pairing — even if the exact two-word \
   combination only appears split across nearby sentences:
   - "Python automation" — when the JD describes Python being used for \
     automation work.
   - "LLM integration" — when the JD describes integrating LLMs with APIs, \
     tools, or systems.
   - "workflow automation" — when the JD describes automating workflows.
   - "API development" — when the JD describes building/designing APIs.
   - "data pipelines" / "ETL pipelines" — when the JD describes pipeline work.
   This is the ONE narrow exception to rule 1's "verbatim" requirement: \
   these standard recruiter-search phrases may be assembled from adjacent \
   JD content, but only when the underlying activity is plainly there.

6. STRIP FILLER. Drop articles, prepositions, generic puffery ("team \
   player", "self-starter", "fast-paced", "5+ years", "great opportunity"), \
   and the company's own product/office names. Keep the substantive head: \
   "5+ years of distributed systems experience" → "distributed systems".

7. KEEP SOFT SKILLS MINIMAL. ATS systems weight soft skills lightly and \
   recruiters glaze over them. Include at most 1-2 soft skills total, and \
   only when explicitly emphasized in the JD. Avoid:
   - "communication skills", "collaboration skills", "work independently", \
     "team player", "attention to detail", "problem solver".
   The exception is when a soft skill is part of a domain-specific phrase \
   ("technical communication", "stakeholder management") — those count as \
   responsibilities, not generic soft skills.

8. DEDUPE AND REMOVE OVERLAP.
   - Casing duplicates → one entry. "Python", "python", "Python 3" → "Python".
   - Generic-vs-specific overlap → keep the specific one. "performance" + \
     "performance optimization" → keep only "performance optimization". \
     "API" + "REST API design" → keep only "REST API design".
   - Exception: keep both an acronym AND its spelled-out form if the JD \
     uses both ("ML" + "Machine Learning").

9. LENGTH: 1-4 words per keyword. Five words is allowed only when no \
   shorter form preserves the meaning. If you've written a 5-word keyword, \
   ask whether 3-4 words would do.

10. CATEGORIZE:
    - skill          : languages, named soft skills (Python, communication)
    - tool           : named products (AWS Glue, React, Snowflake, OpenAI, LangChain, Hugging Face)
    - concept        : methodologies / techniques (ETL, agile, prompt engineering, embeddings, RAG, fine-tuning)
    - responsibility : duties and action+skill phrases (code review, build automation solutions, LLM integration)
    - qualification  : credentials / education (Bachelor's degree, AWS certified)

11. IMPORTANCE:
    - "must" : under "Required" / "Must have" / "Essential", or stated with \
      strong language ("required", "minimum", "proficient in"). Default to \
      "must" when the section is unclear or when the term is in the main \
      body of responsibilities/requirements.
    - "nice" : under "Preferred" / "Nice to have" / "Bonus" / "Plus".

JOB DESCRIPTION (this is the only source — do not look beyond it):
\"\"\"
{jd_text}
\"\"\"
"""

async def extract_jd_keywords(jd_text: str) -> List[JDKeyword]:
    """Extract keywords from the JD using Gemini with a strict no-hallucination
    prompt. Returns [] when Gemini is not configured.

    """
    if not gemini_client or not jd_text.strip():
        return []
    prompt = JD_KEYWORD_PROMPT.format(jd_text=jd_text)
    try:
        response = await gemini_client.aio.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=JDKeywordList,
                temperature=0.1,  # low — we want determinism, not creativity
            ),
        )
        parsed = response.parsed
        if isinstance(parsed, JDKeywordList):
            kws = parsed.keywords
        elif response.text:
            kws = JDKeywordList.model_validate_json(response.text).keywords
        else:
            return []
    except Exception as e:
        print(f"[gemini] JD keyword extraction failed: {e}")
        raise

    # Defensive client-side filter: drop empty terms and dedupe case-insensitively.
    seen: Dict[str, JDKeyword] = {}
    for k in kws:
        term = (k.term or "").strip()
        if not term:
            continue
        key = term.lower()
        if key in seen:
            # Prefer "must" over "nice" if both seen for same term.
            if seen[key].importance == "nice" and k.importance == "must":
                seen[key] = k
            continue
        seen[key] = k
    return list(seen.values())





def _flatten_cv_text(cv: CVData) -> str:
    """Concatenate every user-visible CV field into a single searchable string."""
    parts: List[str] = []
    p = cv.personal
    parts += [p.title, p.location]  # name etc. usually irrelevant for keyword match
    if cv.summary:
        parts.append(cv.summary)
    for e in cv.experience:
        parts += [e.role, e.company, e.location] + list(e.bullets)
    for proj in cv.projects:
        parts += [proj.name, proj.tech, proj.link] + list(proj.bullets)
    for ed in cv.education:
        parts += [ed.school, ed.degree, ed.coursework, ed.gpa, ed.details]
    for s in cv.skills:
        parts += [s.category, s.items]
    for c in cv.certifications:
        parts += [c.name, c.issuer]
    return "\n".join(p for p in parts if p)


def score_cv_against_keywords(
    cv: CVData,
    keywords: List[JDKeyword],
) -> Dict[str, Any]:
    """Compute the ATS keyword-match score for `cv` vs `keywords`.

    Score = matched_count / total_count * 100 (flat, unweighted).
    A keyword matches ONLY if its exact phrase appears in the CV text, with
    word boundaries and case-insensitively. No lemmatization, no token-spread
    matching — what you see in the preview is what counts.
    """
    if not keywords:
        return {
            "score": 0,
            "matched": 0,
            "total": 0,
            "matched_terms": [],
            "missing_terms": [],
        }

    cv_text = _flatten_cv_text(cv).lower()

    matched: List[Dict[str, Any]] = []
    missing: List[Dict[str, Any]] = []

    for kw in keywords:
        term = (kw.term or "").strip()
        if not term:
            continue

        # Word-boundary regex (case-insensitive). Custom boundary instead of \b
        # so terms like "C++", ".NET", "CI/CD" still match correctly — \b
        # treats "+", ".", and "/" as word boundaries which breaks those.
        pattern = (
            r"(?<![A-Za-z0-9])"
            + re.escape(term.lower())
            + r"(?![A-Za-z0-9])"
        )
        hit = re.search(pattern, cv_text) is not None

        entry = {"term": term, "category": kw.category, "importance": kw.importance}
        (matched if hit else missing).append(entry)

    total = len(matched) + len(missing)
    score = round((len(matched) / total) * 100) if total else 0
    return {
        "score": score,
        "matched": len(matched),
        "total": total,
        "matched_terms": matched,
        "missing_terms": missing,
    }


# =====================================================================
# Tailoring (Gemini, evidence-only)
# =====================================================================

TAILOR_PROMPT = """You are an expert resume-tailoring assistant. Your single \
most important rule is HONESTY: you may NEVER invent skills, experiences, \
metrics, employers, projects, or qualifications that are not already supported \
by evidence in the candidate's CV.

You will receive:
1. A candidate's CV (JSON).
2. A list of keywords/phrases extracted from a target job description (JD).

Your job is to produce a list of EDITS to the CV that re-phrase existing \
content using the JD's vocabulary, ONLY when the underlying fact is already \
present in the CV.

ABSOLUTE RULES:
1. EVIDENCE-ONLY. For every JD keyword you weave into the CV, the CV must \
   already contain a fact that supports it. Examples:
   - JD says "CI/CD pipelines"; CV bullet says "automated deployments using \
     CircleCI" → ALLOWED. Rephrase to "Built CI/CD pipelines using CircleCI \
     for automated deployments."
   - JD wants "Kubernetes"; nothing in the CV mentions container orchestration \
     → NOT ALLOWED. Add it to `gaps` instead.
   - JD wants "leadership"; CV says "led a team of 4" → ALLOWED.
2. Never fabricate metrics, dates, employers, project names, or technologies.
3. Preserve every fact: numbers, percentages, dates, employers, scope.
4. Keep bullets the same length or slightly shorter — never pad.
5. Do not change the personal section, education, or certifications. Only \
   touch experience bullets, project bullets, project tech fields, project \
   names (rarely), and skills items.
6. For skills: you MAY add a JD keyword to an existing skills `items` string \
   ONLY IF a synonym or closely related tool is already there. Example: \
   "AWS Glue, Lambda, S3" supports adding "ETL on AWS" to a related concept \
   bullet, but does NOT support adding the standalone tool "Databricks".
7. For each edit, list the JD keywords you wove in under `inserted_terms`. \
   These will be highlighted for the user to review.
8. Anything from the JD keyword list that is NOT supported by the CV: put it \
   in `gaps` with the same category and importance. Be honest — the user \
   needs to know what to add manually.
9. `summary`: one short sentence describing what you changed (e.g. \
   "Rephrased 4 bullets to use JD vocabulary; 3 keywords gap-flagged.").

EDIT PATH FORMAT:
Use dot-separated paths into the CV JSON:
- "experience.<i>.bullets.<j>"   -> a single bullet inside experience[i].bullets[j]
- "projects.<i>.bullets.<j>"     -> a project bullet
- "projects.<i>.tech"            -> the tech string of a project
- "skills.<i>.items"             -> the items string of a skill category
- "summary"                      -> the top-level summary
Indexes are 0-based and must match the input CV exactly.

CV (JSON):
\"\"\"
{cv_json}
\"\"\"

JD KEYWORDS (JSON):
\"\"\"
{kw_json}
\"\"\"
"""


async def tailor_cv_with_gemini(
    cv: CVData,
    keywords: List[JDKeyword],
) -> Optional[TailorResult]:
    if not gemini_client:
        return None
    prompt = TAILOR_PROMPT.format(
        cv_json=json.dumps(cv.model_dump(), indent=2),
        kw_json=json.dumps([k.model_dump() for k in keywords], indent=2),
    )
    try:
        response = await gemini_client.aio.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=TailorResult,
                temperature=0.2,
            ),
        )
        parsed = response.parsed
        if isinstance(parsed, TailorResult):
            return parsed
        if response.text:
            return TailorResult.model_validate_json(response.text)
    except Exception as e:
        print(f"[gemini] tailor failed: {e}")
        raise
    return None


# ----- Apply edits to CV + compute highlight spans --------------------

def _set_path(cv_dict: Dict[str, Any], path: str, value: str) -> bool:
    """Write `value` to `cv_dict` at the dotted `path`. Returns True on success."""
    if not path:
        return False
    parts = path.split(".")
    obj: Any = cv_dict
    for i, part in enumerate(parts[:-1]):
        if part.isdigit():
            idx = int(part)
            if not isinstance(obj, list) or idx >= len(obj):
                return False
            obj = obj[idx]
        else:
            if not isinstance(obj, dict) or part not in obj:
                return False
            obj = obj[part]
    last = parts[-1]
    try:
        if last.isdigit():
            idx = int(last)
            if not isinstance(obj, list) or idx >= len(obj):
                return False
            obj[idx] = value
        else:
            if not isinstance(obj, dict):
                return False
            obj[last] = value
        return True
    except Exception:
        return False


def _compute_spans(text: str, terms: List[str]) -> List[List[int]]:
    """Find case-insensitive occurrences of each `term` in `text` and return
    [[start, end], ...] character spans, merged + sorted, no overlaps.
    """
    if not text or not terms:
        return []
    spans: List[Tuple[int, int]] = []
    lower = text.lower()
    for term in terms:
        t = (term or "").strip()
        if not t:
            continue
        # Word-boundary match where possible; fall back to plain substring for
        # terms containing punctuation that \b doesn't handle well (e.g. "C++").
        pattern = (
            r"(?<![A-Za-z0-9])" + re.escape(t.lower()) + r"(?![A-Za-z0-9])"
        )
        for m in re.finditer(pattern, lower):
            spans.append((m.start(), m.end()))
    if not spans:
        return []
    # Merge overlapping/adjacent spans
    spans.sort()
    merged: List[List[int]] = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return merged


def apply_tailor_result(
    cv: CVData,
    result: TailorResult,
) -> Tuple[CVData, Dict[str, List[List[int]]]]:
    """Apply Gemini's edits to a copy of `cv`. Returns (new_cv, highlights)."""
    cv_dict = cv.model_dump()
    highlights: Dict[str, List[List[int]]] = {}
    for edit in result.edits:
        if not edit.path or edit.new_text is None:
            continue
        if _set_path(cv_dict, edit.path, edit.new_text):
            spans = _compute_spans(edit.new_text, edit.inserted_terms)
            if spans:
                highlights[edit.path] = spans
    new_cv = CVData.model_validate(cv_dict)
    return new_cv, highlights


# =====================================================================
# In-memory session store (swap with Redis / DB in production)
# =====================================================================

class Session:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.created_at = datetime.utcnow().isoformat()
        self.updated_at = self.created_at
        self.uploaded_resume: Optional[Dict[str, Any]] = None
        self.job_description: Optional[Dict[str, Any]] = None
        self.cv_data: Optional[CVData] = None
        # JD keyword extraction (filled by /api/job-description)
        self.jd_keywords: List[Dict[str, Any]] = []
        # Gaps from the most recent /api/build-tailored-cv call.
        # The tailored CV itself is held by the frontend (cvTailored), since
        # the user can edit it independently of the original.
        self.gaps: List[Dict[str, Any]] = []

    def touch(self):
        self.updated_at = datetime.utcnow().isoformat()

    def to_dict(self):
        return {
            "session_id": self.session_id,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
            "uploaded_resume": self.uploaded_resume,
            "job_description": self.job_description,
            "cv_data": self.cv_data.model_dump() if self.cv_data else None,
            "jd_keywords": self.jd_keywords,
            "gaps": self.gaps,
        }


sessions: Dict[str, Session] = {}
DEFAULT_SESSION_ID = "default"
sessions[DEFAULT_SESSION_ID] = Session(DEFAULT_SESSION_ID)


def get_or_create_session(session_id: str) -> Session:
    if session_id not in sessions:
        sessions[session_id] = Session(session_id)
    return sessions[session_id]


# =====================================================================
# Text extraction helpers
# =====================================================================

def extract_text_from_pdf(content: bytes) -> str:
    if not HAS_PYPDF:
        raise HTTPException(500, "pypdf not installed; run `pip install pypdf`")
    reader = pypdf.PdfReader(io.BytesIO(content))
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts).strip()


def extract_text_from_docx(content: bytes) -> str:
    if not HAS_DOCX:
        raise HTTPException(500, "python-docx not installed; run `pip install python-docx`")
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p for p in parts if p.strip()).strip()


def extract_text_from_txt(content: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()


def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(content)
    if name.endswith(".docx"):
        return extract_text_from_docx(content)
    if name.endswith(".txt") or name.endswith(".md"):
        return extract_text_from_txt(content)
    try:
        return extract_text_from_txt(content)
    except Exception:
        raise HTTPException(415, f"Unsupported file type: {filename}")


# =====================================================================
# Endpoints
# =====================================================================

@app.get("/")
def index():
    return FileResponse(_static_dir() / "resume-builder.html")

@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "time": datetime.utcnow().isoformat(),
        "active_sessions": len(sessions),
        "extractors": {"pdf": HAS_PYPDF, "docx": HAS_DOCX, "txt": True},
        "gemini": {
            "sdk_installed": HAS_GENAI,
            "configured": gemini_client is not None,
            "model": GEMINI_MODEL if gemini_client else None,
        },
    }


@app.post("/api/session/new")
def new_session():
    sid = str(uuid.uuid4())
    sessions[sid] = Session(sid)
    return {"session_id": sid, "created_at": sessions[sid].created_at}


@app.get("/api/session/{session_id}")
def get_session(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    return sessions[session_id].to_dict()


@app.delete("/api/session/{session_id}")
def reset_session(session_id: str):
    sessions[session_id] = Session(session_id)
    return {"reset": True, "session_id": session_id}


@app.post("/api/upload-resume")
async def upload_resume(
    file: UploadFile = File(...),
    session_id: str = Form(DEFAULT_SESSION_ID),
    auto_parse: str = Form("true"),
):
    """
    Accept a past resume (PDF/DOCX/TXT), extract its plain text,
    auto-parse it with Gemini into CVData, and store in the session.
    """
    content = await file.read()
    if not content:
        raise HTTPException(400, "Empty file")
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(413, "File too large (max 10 MB)")

    text = extract_text(file.filename or "", content)

    sess = get_or_create_session(session_id)
    sess.uploaded_resume = {
        "filename": file.filename,
        "content_type": file.content_type,
        "size_bytes": len(content),
        "char_count": len(text),
        "text": text,
        "uploaded_at": datetime.utcnow().isoformat(),
    }

    parsed_cv: Optional[CVData] = None
    parse_error: Optional[str] = None

    if auto_parse.lower() in ("true", "1", "yes") and gemini_client and text:
        try:
            parsed_cv = await parse_resume_with_gemini(text)
            if parsed_cv:
                sess.cv_data = parsed_cv
        except Exception as e:
            parse_error = str(e)

    sess.touch()

    return {
        "session_id": session_id,
        "filename": file.filename,
        "size_bytes": len(content),
        "char_count": len(text),
        "preview": text[:500] + ("…" if len(text) > 500 else ""),
        "gemini_available": gemini_client is not None,
        "parsed_cv": parsed_cv.model_dump() if parsed_cv else None,
        "parse_error": parse_error,
    }


@app.post("/api/parse-resume/{session_id}")
async def parse_existing_resume(session_id: str):
    """Re-run Gemini parsing on the resume already uploaded in this session."""
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    sess = sessions[session_id]
    if not sess.uploaded_resume:
        raise HTTPException(400, "No resume uploaded in this session")
    if not gemini_client:
        raise HTTPException(
            503,
            "Gemini not configured. Set GEMINI_API_KEY in environment or .env file.",
        )

    text = sess.uploaded_resume.get("text", "")
    try:
        parsed = await parse_resume_with_gemini(text)
    except Exception as e:
        raise HTTPException(500, f"Gemini parse failed: {e}")

    if not parsed:
        raise HTTPException(500, "Gemini returned no parseable result")

    sess.cv_data = parsed
    sess.touch()
    return {"data": parsed.model_dump()}


@app.post("/api/job-description")
async def save_job_description(payload: JDPayload):
    """Save the target job description AND extract keywords with Gemini.

    Uses a strict no-hallucination prompt — every returned keyword is one
    that literally appears in the JD text. Stores the keyword list on the
    session and computes an initial score against the current CV (if any).
    """
    sess = get_or_create_session(payload.session_id)
    sess.job_description = {
        "text": payload.text,
        "char_count": len(payload.text),
        "saved_at": datetime.utcnow().isoformat(),
    }

    keywords: List[JDKeyword] = []
    extract_error: Optional[str] = None

    if not gemini_client:
        extract_error = "Gemini not configured. Set GEMINI_API_KEY in environment or .env file."
        sess.jd_keywords = []
    elif payload.text.strip():
        try:
            keywords = await extract_jd_keywords(payload.text)
            sess.jd_keywords = [k.model_dump() for k in keywords]
        except Exception as e:
            extract_error = str(e)
            sess.jd_keywords = []
    else:
        sess.jd_keywords = []

    # Initial score against the current CV (if any)
    score_result: Dict[str, Any] = {
        "score": 0, "matched": 0, "total": len(keywords),
        "matched_terms": [], "missing_terms": [],
    }
    if sess.cv_data and keywords:
        score_result = score_cv_against_keywords(sess.cv_data, keywords)

    sess.touch()
    return {
        "session_id": payload.session_id,
        "char_count": len(payload.text),
        "saved_at": sess.job_description["saved_at"],
        "keywords": sess.jd_keywords,
        "keyword_count": len(sess.jd_keywords),
        "gemini_available": gemini_client is not None,
        "extract_error": extract_error,
        "score": score_result,
    }


@app.post("/api/cv-data")
def save_cv_data(payload: CVPayload):
    """Persist the current state of the CV form (matching cv_template.json)."""
    sess = get_or_create_session(payload.session_id)
    sess.cv_data = payload.data
    sess.touch()
    return {"session_id": payload.session_id, "saved": True}


@app.get("/api/cv-data/{session_id}")
def get_cv_data(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    sess = sessions[session_id]
    return {"data": sess.cv_data.model_dump() if sess.cv_data else None}


# ----- New: live ATS scoring -----------------------------------------

@app.post("/api/score")
def score_cv(payload: ScorePayload):
    """Score `payload.data` (or the session's CV if not provided) against the
    keywords currently stored on the session. 
    """
    sess = get_or_create_session(payload.session_id)
    cv = payload.data or sess.cv_data
    if cv is None:
        return {
            "score": 0, "matched": 0, "total": len(sess.jd_keywords),
            "matched_terms": [], "missing_terms": sess.jd_keywords,
            "no_cv": True,
        }
    keywords = [JDKeyword(**k) for k in sess.jd_keywords]
    return score_cv_against_keywords(cv, keywords)


@app.post("/api/jd-keywords")
def update_jd_keywords(payload: UpdateKeywordsPayload):
    """Replace the session's keyword list with `payload.keywords` and
    re-score against the original CV (and optionally a tailored CV passed
    in the body). Used when the user removes a keyword chip from the panel.

    No LLM call — the keywords come from the client, fully under user control.
    """
    sess = get_or_create_session(payload.session_id)

    # Normalize: drop empties, dedupe by lowercased term, keep first occurrence.
    seen = set()
    cleaned: List[JDKeyword] = []
    for k in payload.keywords:
        term = (k.term or "").strip()
        if not term:
            continue
        key = term.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(k)

    sess.jd_keywords = [k.model_dump() for k in cleaned]
    sess.touch()

    # Re-score the original (always) and the tailored CV (when sent).
    score_original = (
        score_cv_against_keywords(sess.cv_data, cleaned)
        if sess.cv_data and cleaned
        else {"score": 0, "matched": 0, "total": len(cleaned),
              "matched_terms": [], "missing_terms": [k.model_dump() for k in cleaned]}
    )
    score_tailored = None
    if payload.tailored_data and cleaned:
        score_tailored = score_cv_against_keywords(payload.tailored_data, cleaned)

    return {
        "session_id":     payload.session_id,
        "keywords":       sess.jd_keywords,
        "keyword_count":  len(sess.jd_keywords),
        "score_original": score_original,
        "score_tailored": score_tailored,
    }


# ----- Build the ATS-friendly tailored CV ----------------------------
# This does NOT overwrite the session's CV — the frontend keeps both versions
# (original + tailored) and lets the user toggle between them. We just send
# back the new CV plus its score. Gemini is restricted to evidence-only
# rephrasing; missing JD requirements come back in `gaps`.

@app.post("/api/build-tailored-cv")
async def build_tailored_cv(payload: TailorPayload):
    sess = sessions.get(payload.session_id)
    if not sess:
        raise HTTPException(404, "Session not found")
    if not sess.cv_data:
        raise HTTPException(400, "No CV data on this session — upload or sync a CV first.")
    if not sess.jd_keywords:
        raise HTTPException(
            400,
            "No JD keywords on this session — save a job description first.",
        )
    if not gemini_client:
        raise HTTPException(
            503,
            "Gemini not configured. Set GEMINI_API_KEY in environment or .env file.",
        )

    keywords = [JDKeyword(**k) for k in sess.jd_keywords]

    try:
        result = await tailor_cv_with_gemini(sess.cv_data, keywords)
    except Exception as e:
        raise HTTPException(500, f"Gemini tailor failed: {e}")
    if not result:
        raise HTTPException(500, "Gemini returned no tailor result")

    # Apply edits to a COPY of the CV — original stays untouched on the
    # session. The frontend stores the returned CV as `cvTailored`.
    tailored_cv, _ = apply_tailor_result(sess.cv_data, result)
    sess.gaps = [g.model_dump() for g in result.gaps]
    sess.touch()

    # Score both versions so the UI can show "Original X% → Tailored Y%".
    score_original = score_cv_against_keywords(sess.cv_data, keywords)
    score_tailored = score_cv_against_keywords(tailored_cv, keywords)

    return {
        "tailored_cv":   tailored_cv.model_dump(),
        "gaps":          sess.gaps,
        "summary":       result.summary,
        "edits_applied": len(result.edits),
        "score_original": score_original,
        "score_tailored": score_tailored,
    }


# =====================================================================
# Run with: uvicorn main:app --reload --port 8000
# =====================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, port=8000)
